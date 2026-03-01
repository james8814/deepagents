"""Word document converter using python-docx."""

import logging
import time
from pathlib import Path

from deepagents.middleware.converters.base import BaseConverter

logger = logging.getLogger(__name__)


class DOCXConverter(BaseConverter):
    """Converter for Word documents to Markdown.

    Supports .docx files with:
    - Text and paragraphs
    - Headings and styles
    - Tables
    - Lists

    Dependencies:
        python-docx: Install with `pip install python-docx`
    """

    def convert(self, path: Path, raw_content: str | bytes | None = None) -> str:
        """Convert a Word document to Markdown.

        Args:
            path: Path to the Word document.
            raw_content: Ignored for Word documents (binary format).

        Returns:
            Markdown-formatted string.
        """
        from docx import Document

        start = time.time()

        doc = Document(path)
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check style for heading level
            style_name = para.style.name.lower() if para.style else ""

            if "heading 1" in style_name:
                parts.append(f"# {text}\n")
            elif "heading 2" in style_name:
                parts.append(f"## {text}\n")
            elif "heading 3" in style_name:
                parts.append(f"### {text}\n")
            elif "heading 4" in style_name:
                parts.append(f"#### {text}\n")
            elif "heading 5" in style_name:
                parts.append(f"##### {text}\n")
            elif "heading 6" in style_name:
                parts.append(f"###### {text}\n")
            elif "list" in style_name:
                # Handle list items
                parts.append(f"- {text}")
            else:
                parts.append(text)

        # Extract tables
        for i, table in enumerate(doc.tables, start=1):
            parts.append(f"\n### Table {i}\n")

            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(row_data)

            if rows:
                headers = rows[0]
                data_rows = rows[1:] if len(rows) > 1 else []
                parts.append(self._format_as_table(data_rows, headers))

        duration = time.time() - start
        self._log_conversion(path, duration)

        return "\n\n".join(parts)